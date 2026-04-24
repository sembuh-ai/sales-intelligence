#!/usr/bin/env python3
"""
Sales Intelligence — Unified Automation
Combines Monday.com CRM operations with document generation pipeline.

Commands:
    python app.py generate              — Fetch deals → generate docs → upload → email → Slack
    python app.py generate --deal "X"   — Generate docs for one deal only
    python app.py mom <file.docx>       — Extract MoM and update Monday.com CRM
    python app.py mom <file.docx> --dry-run
    python app.py health                — Run deal health scoring
    python app.py interactive           — Interactive Monday.com assistant
"""

import asyncio
import argparse
import base64
import copy
import datetime
import json
import os
import re
import sqlite3
import sys
from contextlib import AsyncExitStack

import requests
from anthropic import Anthropic
from dotenv import load_dotenv

# ── Load env from all sources ─────────────────────────────────
_BASE_DIR = os.path.dirname(os.path.abspath(__file__))
load_dotenv(os.path.join(_BASE_DIR, ".env"))
load_dotenv(os.path.join(_BASE_DIR, "monday", ".env"), override=False)
load_dotenv(os.path.join(_BASE_DIR, "gmcp", ".env"), override=False)

# ── Config ────────────────────────────────────────────────────
MONDAY_API_KEY = os.getenv("MONDAY_API_KEY", "")
MONDAY_WORKSPACE_ID = os.getenv("MONDAY_WORKSPACE_ID", "")
MONDAY_WORKSPACE_NAME = os.getenv("MONDAY_WORKSPACE_NAME", "")
CLAUDE_MODEL = os.getenv("CLAUDE_MODEL", "claude-opus-4-6")
SLACK_BOT_TOKEN = os.getenv("SLACK_BOT_TOKEN", "")
SLACK_CHANNEL_BOD = "C0AUZ7GB3DJ"   # #bod-updates
SLACK_CHANNEL_AM = "C0AUE6XG8TZ"    # #am-indi
SLACK_NOTIFY_USER = "U08UV06KD45"
GOOGLE_FOLDER_ID = os.getenv("FOLDER_ID", "")

MONDAY_API_URL = "https://api.monday.com/v2"
SLACK_API_URL = "https://slack.com/api"
DB_PATH = os.path.join(_BASE_DIR, "sales_intelligence.db")

TEMPLATE_DIR = os.path.join(_BASE_DIR, "PTP Hackathon - Brief")
OUTPUT_DIR = os.path.join(_BASE_DIR, "output")

TODAY = datetime.date.today()
TODAY_STR = TODAY.strftime("%d %B %Y")

# ── Pricing / Product Config ─────────────────────────────────
MARGIN_ASSUMPTIONS = {
    "Claim Workflow": {"unit_price": 1.0, "margin": 0.60, "type": "Recurring", "uom": "Per Claim"},
    "Fraud Detection": {"unit_price": 2.0, "margin": 0.65, "type": "Recurring", "uom": "Per Claim"},
    "Implementation": {"unit_price": 50000, "margin": 0.50, "type": "One Time Fee", "uom": "Per Implementation"},
}

STAGE_WEIGHTS = {
    "Open": 0.10, "First Meeting Done": 0.20, "Piloting": 0.30,
    "Solutioning": 0.50, "Waiting Confirmation": 0.70,
    "[Won] Waiting Signature": 0.90, "Won": 1.00,
    "Implementation Done": 1.00, "Lost": 0.00,
}

HEALTH_RULES = [
    # (rule_name, deduction, check_fn)
    # check_fn(deal, today) -> bool (True = deduction applies)
]


def detect_products(deal_name):
    name_lower = deal_name.lower()
    if "full suite" in name_lower:
        return ["Claim Workflow", "Fraud Detection", "Implementation"]
    if "fwa" in name_lower or "fraud" in name_lower:
        return ["Fraud Detection", "Implementation"]
    if "ocr" in name_lower or "stp" in name_lower:
        return ["Claim Workflow", "Implementation"]
    return ["Claim Workflow", "Fraud Detection", "Implementation"]


def _parse_num(val):
    if not val:
        return 0
    val = str(val).replace(",", "").replace("$", "").strip()
    try:
        return float(val)
    except ValueError:
        return 0


# ══════════════════════════════════════════════════════════════
# MONDAY.COM API (Direct GraphQL)
# ══════════════════════════════════════════════════════════════

def monday_query(query, variables=None):
    headers = {
        "Authorization": MONDAY_API_KEY,
        "Content-Type": "application/json",
        "API-Version": "2024-10",
    }
    payload = {"query": query}
    if variables:
        payload["variables"] = variables
    resp = requests.post(MONDAY_API_URL, json=payload, headers=headers)
    resp.raise_for_status()
    data = resp.json()
    if "errors" in data:
        raise RuntimeError(f"Monday API error: {data['errors']}")
    return data["data"]


def fetch_boards():
    query = """
    query ($wsIds: [ID!]) {
        boards(workspace_ids: $wsIds, limit: 50) { id name }
    }
    """
    return monday_query(query, {"wsIds": [int(MONDAY_WORKSPACE_ID)]})["boards"]


def find_deals_board(boards):
    for b in boards:
        if b["name"].strip().lower() == "deals":
            return b
    for b in boards:
        if "deal" in b["name"].lower() and "subitem" not in b["name"].lower():
            return b
    return boards[0] if boards else None


def fetch_deals(board_id):
    query = """
    query ($boardId: [ID!]!) {
        boards(ids: $boardId) {
            items_page(limit: 100) {
                items {
                    id name updated_at
                    column_values { id type text value }
                }
            }
        }
    }
    """
    data = monday_query(query, {"boardId": [int(board_id)]})
    items = data["boards"][0]["items_page"]["items"]
    deals = []
    for item in items:
        deal = {"id": item["id"], "name": item["name"], "updated_at": item.get("updated_at", "")}
        for col in item["column_values"]:
            deal[col["id"]] = col.get("text", "") or ""
            deal[f"{col['id']}_raw"] = col.get("value", "")
        deals.append(deal)
    return deals


def parse_deal(deal):
    name = deal.get("name", "Unknown")
    stage = deal.get("deal_stage", "Open")
    deal_value = _parse_num(deal.get("deal_value", "0"))
    close_date = deal.get("deal_expected_close_date", "")
    monthly_claim_vol = _parse_num(deal.get("numeric_mm1bmx9t", "0"))
    members_covered = _parse_num(deal.get("numeric_mm1bx91m", "0"))
    pricing_model = deal.get("dropdown_mm1b79r5", "Per Claim")
    owner = deal.get("deal_owner", "")
    proposal_date = deal.get("date_mm1bpvvx", "")
    incurred = _parse_num(deal.get("numeric_mm1bdpzy", "0"))
    excess = _parse_num(deal.get("numeric_mm1bkxy8", "0"))
    approved = _parse_num(deal.get("numeric_mm1b64b7", "0"))
    # Last interaction date (date__1) or fall back to Monday item updated_at
    last_contact = deal.get("date__1", "") or ""
    updated_at = deal.get("updated_at", "")

    products = detect_products(name)

    # Estimate volume from deal_value when not set
    if monthly_claim_vol == 0 and deal_value > 0:
        impl_fee = 50000 if "Implementation" in products else 0
        recurring_rate = sum(
            MARGIN_ASSUMPTIONS[p]["unit_price"] for p in products if p != "Implementation"
        ) or 1
        annual_claim_vol = max(int((deal_value - impl_fee) / recurring_rate), 1200) if deal_value > impl_fee else 12000
        monthly_claim_vol = annual_claim_vol / 12
    else:
        annual_claim_vol = monthly_claim_vol * 12

    if members_covered == 0 and monthly_claim_vol > 0:
        members_covered = monthly_claim_vol * 10

    return {
        "id": deal.get("id"), "name": name, "stage": stage,
        "deal_value": deal_value, "close_date": close_date,
        "monthly_claim_vol": monthly_claim_vol, "annual_claim_vol": annual_claim_vol,
        "members_covered": members_covered, "pricing_model": pricing_model or "Per Claim",
        "incurred": incurred, "excess": excess, "approved": approved,
        "owner": owner, "proposal_date": proposal_date, "products": products,
        "last_contact": last_contact, "updated_at": updated_at,
    }


# ══════════════════════════════════════════════════════════════
# MONDAY.COM MCP CLIENT (for MoM + Interactive)
# ══════════════════════════════════════════════════════════════

class MondayMCPClient:
    """Claude-powered Monday.com MCP client for complex operations."""

    def __init__(self):
        self.exit_stack = AsyncExitStack()
        self.session = None
        self.anthropic = Anthropic()
        self.tools = []

    async def connect(self):
        from mcp import ClientSession, StdioServerParameters
        from mcp.client.stdio import stdio_client

        server_params = StdioServerParameters(
            command="npx",
            args=["-y", "verdant-monday-mcp"],
            env={
                **os.environ,
                "MONDAY_API_KEY": MONDAY_API_KEY,
                "MONDAY_WORKSPACE_NAME": MONDAY_WORKSPACE_NAME,
            },
        )
        transport = await self.exit_stack.enter_async_context(stdio_client(server_params))
        read, write = transport
        self.session = await self.exit_stack.enter_async_context(
            __import__("mcp").ClientSession(read, write)
        )
        await self.session.initialize()

        response = await self.session.list_tools()
        self.tools = [
            {"name": t.name, "description": t.description or "", "input_schema": t.inputSchema}
            for t in response.tools
        ]
        print(f"MCP connected. {len(self.tools)} Monday tools.", file=sys.stderr)

    async def chat(self, prompt: str, history: list, system: str = "") -> tuple[str, list]:
        history.append({"role": "user", "content": prompt})

        default_system = (
            f"You are a Monday.com CRM assistant for workspace {MONDAY_WORKSPACE_ID}. "
            f"Always pass workspace_ids=[{MONDAY_WORKSPACE_ID}] when listing boards. "
            "Be concise. Use tables/lists where helpful."
        )

        while True:
            response = self.anthropic.messages.create(
                model=CLAUDE_MODEL, max_tokens=4096,
                system=system or default_system,
                messages=history, tools=self.tools,
            )
            assistant_content = list(response.content)
            history.append({"role": "assistant", "content": assistant_content})

            if response.stop_reason == "end_turn":
                return "".join(b.text for b in response.content if b.type == "text"), history

            tool_results = []
            for block in response.content:
                if block.type != "tool_use":
                    continue
                print(f"  calling {block.name}...", file=sys.stderr)

                tool_input = block.input
                if block.name == "monday_create_update" and "updateText" in tool_input:
                    tool_input = {**tool_input, "updateText": tool_input["updateText"].replace("\n", "<br>")}

                try:
                    result = await self.session.call_tool(block.name, tool_input)
                    result_text = " ".join(c.text for c in (result.content or []) if hasattr(c, "text"))
                    tool_results.append({"type": "tool_result", "tool_use_id": block.id, "content": result_text})
                except Exception as e:
                    print(f"  TOOL ERROR: {e}", file=sys.stderr)
                    tool_results.append({"type": "tool_result", "tool_use_id": block.id, "content": f"Error: {e}", "is_error": True})

            history.append({"role": "user", "content": tool_results})

    async def cleanup(self):
        await self.exit_stack.aclose()


# ══════════════════════════════════════════════════════════════
# DOCUMENT GENERATORS
# ══════════════════════════════════════════════════════════════

def _unmerge_and_clear(ws):
    for merge in list(ws.merged_cells.ranges):
        ws.unmerge_cells(str(merge))
    for row in ws.iter_rows(min_row=1, max_row=ws.max_row, max_col=ws.max_column):
        for cell in row:
            cell.value = None


def generate_quotation(deal, seq_num=1, discount_pct=0, discount_note="", client_contact_name="", client_contact_title="", mom_context=""):
    """Generate quotation from Quotation-Sentosa-Health-Demo template.
    Yellow cells = dynamic values from Monday CRM and MoM.
    """
    import openpyxl

    template_path = os.path.join(TEMPLATE_DIR, "Quotation-Sentosa-Health-Demo.xlsx")
    wb = openpyxl.load_workbook(template_path)
    ws = wb.active

    client_name = deal["name"]
    q_number = f"SBH-Q-{TODAY.year}-{seq_num:03d}"
    validity_date = (TODAY + datetime.timedelta(days=30)).strftime("%d %B %Y")
    products = deal["products"]
    products_str = ", ".join(products)
    monthly_vol = int(deal["monthly_claim_vol"] or 0)
    annual_vol = int(deal["annual_claim_vol"] or monthly_vol * 12)
    members = int(deal["members_covered"] or monthly_vol * 10)
    fx = 16000
    ip_ratio = 0.30
    op_ratio = 0.70

    # Product monthly IDR costs (from MARGIN_ASSUMPTIONS, converted)
    has_ocr = "Claim Workflow" in products
    has_fwa = "Fraud Detection" in products
    has_stp = "STP" in products or ("Claim Workflow" in products and "Fraud Detection" in products)
    has_impl = "Implementation" in products

    ocr_monthly = MARGIN_ASSUMPTIONS["Claim Workflow"]["unit_price"] * monthly_vol * fx if has_ocr else 0
    fwa_monthly = MARGIN_ASSUMPTIONS["Fraud Detection"]["unit_price"] * monthly_vol * fx if has_fwa else 0
    stp_monthly = 0.5 * monthly_vol * fx if has_stp else 0  # ~$0.50/claim for STP
    impl_fee = MARGIN_ASSUMPTIONS["Implementation"]["unit_price"] * fx if has_impl else 35000000

    # Prepared for name
    prepared_for_name = client_contact_name or client_name
    if client_contact_title:
        prepared_for_name = f"{prepared_for_name} | {client_name}"
    else:
        prepared_for_name = f"{prepared_for_name} | {client_name}"

    # ── Yellow cells: Dynamic values ──

    # H4: Quotation number
    ws["H4"] = q_number

    # I5: Date (yellow)
    ws["I5"] = TODAY.strftime("%B %d, %Y")

    # I6: Valid until (yellow)
    ws["I6"] = (TODAY + datetime.timedelta(days=30)).strftime("%B %d, %Y")

    # B7: Prepared For (yellow) — client contact from MoM
    ws["B7"] = f"Prepared For:\n\n{prepared_for_name}\n\n"

    # B10: Project description (yellow) — from MoM context + CRM data
    products_list = []
    if has_ocr:
        products_list.append("OCR")
    if has_fwa:
        products_list.append("FWA")
    if has_stp:
        products_list.append("STP")
    modules_str = " + ".join(products_list) or products_str

    desc = f"AI-Powered Claims Processing Solution — {modules_str} modules for {client_name}.\n"
    desc += f"Scope: {members:,} active members, ~{monthly_vol:,} claims/month (IP {int(ip_ratio*100)}% / OP {int(op_ratio*100)}%)."
    if mom_context:
        # Add PoC reference if available
        if "PoC" in mom_context or "poc" in mom_context.lower():
            desc += f"\nFollowing successful PoC review."
    if discount_pct:
        desc += f"\n{discount_pct}% commercial discount applied per client request."
    ws["B10"] = desc

    # I12: OCR monthly cost (yellow)
    ws["I12"] = ocr_monthly if has_ocr else 0

    # I13: FWA monthly cost (yellow)
    ws["I13"] = fwa_monthly if has_fwa else 0

    # I14: STP monthly cost (yellow)
    ws["I14"] = stp_monthly if has_stp else 0

    # I15: Implementation fee (yellow)
    ws["I15"] = impl_fee

    # I16-I19: formulas already in template (monthly subtotal, annual, impl, gross)
    # I20: Discount row — update formula based on actual discount
    if discount_pct:
        ws["I20"] = f"=I19*-{discount_pct/100}"
        ws["B20"] = f"Discount ({discount_pct}%)"
    else:
        ws["I20"] = 0
        ws["B20"] = "Discount (0%)"

    # I21-I22: formulas already in template (grand total, effective monthly)

    # B23: Discount note (yellow)
    if discount_pct and discount_note:
        ws["B23"] = f"Discount: {discount_pct}% — {discount_note}"
    elif discount_pct:
        ws["B23"] = f"Discount: {discount_pct}% applied per commercial negotiation."
    else:
        ws["B23"] = ""

    # B33: Terms & Conditions — update members/volume references
    terms = ws["B33"].value or ""
    terms = terms.replace("62,000", f"{members:,}")
    terms = terms.replace("4,200", f"{monthly_vol:,}")
    if discount_pct:
        # Update discount reference in T&C
        terms = terms.replace("10% discount", f"{discount_pct}% discount")
    else:
        # Remove discount T&C clause if no discount
        lines = terms.split("\n")
        lines = [l for l in lines if "discount" not in l.lower() or "threshold" in l.lower()]
        terms = "\n".join(lines)
    ws["B33"] = terms

    # ── Sheet 2: ROI & Discount Rules — update dynamic values ──
    if "ROI & Discount Rules" in wb.sheetnames:
        ws2 = wb["ROI & Discount Rules"]
        ws2["H1"] = members  # Member all
        ws2["H2"] = members  # Members active
        ws2["H5"] = annual_vol  # Annual claims
        ws2["H6"] = annual_vol  # Claims all
        incurred = deal.get("incurred", 0)
        if incurred:
            ws2["B4"] = incurred * fx if incurred < 1e9 else incurred  # Yearly incurred IDR
        # Update discount rules — Sentosa Applied row
        if discount_pct:
            ws2["B30"] = f"Applied: {discount_pct}% total discount. {discount_note or 'Per commercial negotiation.'}"

    safe_name = re.sub(r'[^\w\s-]', '', client_name).strip().replace(' ', '_')
    out_path = os.path.join(OUTPUT_DIR, safe_name, f"Sembuh AI - {client_name} - Quotation v1.xlsx")
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    wb.save(out_path)
    print(f"    Quotation: {out_path}" + (f" (discount {discount_pct}%)" if discount_pct else ""))
    return out_path


def generate_pricing_internal(deal, all_deals, discount_pct=0, discount_note=""):
    import openpyxl

    template_path = os.path.join(TEMPLATE_DIR, "Hackathon_Pricing Internal.xlsx")
    wb = openpyxl.load_workbook(template_path)
    client_name = deal["name"]
    products = deal["products"]
    discount_mult = 1 - (discount_pct / 100) if discount_pct else 1.0

    # --- Sheet 1: Per-Client ---
    ws1 = wb.worksheets[0]
    ws1.title = client_name[:31]
    _unmerge_and_clear(ws1)

    ws1["A1"], ws1["H1"] = "Per Client Gross Profit", f"Date {TODAY_STR}"
    hdrs1 = ["Client", "Status", "Total Revenue ($)", "Blended Margin %", "Cost of Revenue ($)", "Gross Profit ($)"]
    if discount_pct:
        hdrs1.extend(["Discount %", "Discount Note"])
    for col, hdr in enumerate(hdrs1):
        ws1.cell(row=2, column=col + 1, value=hdr)

    total_rev = sum(
        MARGIN_ASSUMPTIONS[p]["unit_price"] * discount_mult * (1 if p == "Implementation" else deal["annual_claim_vol"])
        for p in products
    )
    blended_margin = round(sum(MARGIN_ASSUMPTIONS[p]["margin"] for p in products) / len(products), 2) if products else 0.60

    ws1["A3"], ws1["B3"], ws1["C3"], ws1["D3"] = client_name, deal["stage"], round(total_rev, 2), blended_margin
    ws1["E3"], ws1["F3"] = "=C3*(1-D3)", "=C3*D3"
    if discount_pct:
        ws1["G3"] = f"{discount_pct}%"
        ws1["H3"] = discount_note or ""

    # --- Sheet 2: Margin Summary ---
    ws2 = wb.worksheets[1] if len(wb.worksheets) > 1 else wb.create_sheet("Margin Summary")
    _unmerge_and_clear(ws2)

    ws2["A1"] = "Sembuh AI — Gross Profit & Margin Summary"
    ws2["A2"] = "Margin assumptions: Claim Workflow 60% | Fraud Detection 65% | Implementation 50%"
    for col, hdr in enumerate(["Product / Service", "Type", "Unit Price ($)", "Margin %", "Cost per Unit ($)", "Gross Profit per Unit ($)", "Notes"]):
        ws2.cell(row=4, column=col + 1, value=hdr)

    for i, (name, typ, price, margin, notes) in enumerate([
        ("Claim Workflow", "Recurring", 1.0, 0.60, "AI-enabled claim automation"),
        ("Fraud Detection", "Recurring", 2.0, 0.65, "ML inference cost is low"),
        ("Implementation", "One Time Fee", 50000, 0.50, "Engineering + integration"),
    ]):
        r = 5 + i
        ws2[f"A{r}"], ws2[f"B{r}"], ws2[f"C{r}"], ws2[f"D{r}"] = name, typ, price, margin
        ws2[f"E{r}"], ws2[f"F{r}"], ws2[f"G{r}"] = f"=C{r}*(1-D{r})", f"=C{r}*D{r}", notes

    ws2["A10"] = "Per Client Gross Profit"
    for col, hdr in enumerate(["Client", "Status", "Total Revenue ($)", "Blended Margin %", "Cost of Revenue ($)", "Gross Profit ($)", "GP as % of Pipeline"]):
        ws2.cell(row=11, column=col + 1, value=hdr)

    for i, d in enumerate(all_deals):
        r = 12 + i
        rev = sum(MARGIN_ASSUMPTIONS[p]["unit_price"] * (1 if p == "Implementation" else d["annual_claim_vol"]) for p in d["products"])
        bm = round(sum(MARGIN_ASSUMPTIONS[p]["margin"] for p in d["products"]) / len(d["products"]), 2) if d["products"] else 0.6
        ws2[f"A{r}"], ws2[f"B{r}"], ws2[f"C{r}"], ws2[f"D{r}"] = d["name"], d["stage"], round(rev, 2), bm
        ws2[f"E{r}"], ws2[f"F{r}"] = f"=C{r}*(1-D{r})", f"=C{r}*D{r}"

    tr = 12 + len(all_deals)
    ws2[f"A{tr}"] = "TOTAL"
    ws2[f"C{tr}"], ws2[f"D{tr}"] = f"=SUM(C12:C{tr-1})", f"=F{tr}/C{tr}"
    ws2[f"E{tr}"], ws2[f"F{tr}"] = f"=SUM(E12:E{tr-1})", f"=SUM(F12:F{tr-1})"

    # --- Sheet 3: Deal Summary ---
    ws3 = wb.worksheets[2] if len(wb.worksheets) > 2 else wb.create_sheet("Deal Summary")
    _unmerge_and_clear(ws3)
    ws3["A1"] = "Sembuh AI — Potential Client Pipeline"
    ws3["A2"] = "Products & Services Proposal Summary  |  Annual Basis"

    for col, hdr in enumerate(["Client", "Status", "Products", "Monthly Vol", "Annual Vol",
                                "CW $/claim", "FD $/claim", "CW Revenue", "FD Revenue",
                                "Impl Fee", "Total Revenue", "Notes", "Margin %", "Cost", "GP"]):
        ws3.cell(row=4, column=col + 1, value=hdr)

    for i, d in enumerate(all_deals):
        r = 5 + i
        prods = d["products"]
        has_cw, has_fd = "Claim Workflow" in prods, "Fraud Detection" in prods
        bm = round(sum(MARGIN_ASSUMPTIONS[p]["margin"] for p in prods) / len(prods), 2) if prods else 0.6
        ws3[f"A{r}"], ws3[f"B{r}"] = d["name"], d["stage"]
        ws3[f"C{r}"], ws3[f"D{r}"], ws3[f"E{r}"] = " + ".join(prods), d["monthly_claim_vol"], f"=D{r}*12"
        ws3[f"F{r}"] = 1.0 if has_cw else "-"
        ws3[f"G{r}"] = 2.0 if has_fd else "-"
        ws3[f"H{r}"] = f"=E{r}*F{r}" if has_cw else ""
        ws3[f"I{r}"] = f"=E{r}*G{r}" if has_fd else ""
        ws3[f"J{r}"] = 50000 if "Implementation" in prods else ""
        ws3[f"K{r}"], ws3[f"M{r}"] = f"=H{r}+I{r}+J{r}", bm
        ws3[f"N{r}"], ws3[f"O{r}"] = f"=K{r}*(1-M{r})", f"=K{r}*M{r}"

    # --- Sheet 4: Client Detail ---
    ws4 = wb.worksheets[3] if len(wb.worksheets) > 3 else wb.create_sheet("Client Detail")
    _unmerge_and_clear(ws4)
    ws4["A1"] = "Sembuh AI — Per Client Product & Pricing Detail"
    ws4["A2"] = "Annual revenue breakdown per client"
    if discount_pct:
        ws4["A3"] = f"** Discount {discount_pct}% applied to {client_name}: {discount_note or 'per client request'} **"

    for col, hdr in enumerate(["Client", "Product", "Type", "Unit Price ($)", "Discount %", "Discounted Price ($)", "Annual Vol", "UoM", "Revenue ($)", "Notes", "Margin %", "Cost ($)", "GP ($)"]):
        ws4.cell(row=4, column=col + 1, value=hdr)

    detail_notes = {"Claim Workflow": "AI claim workflow", "Fraud Detection": "FWA detection", "Implementation": "Core system integration"}
    row_num = 5
    for d in all_deals:
        is_discounted = discount_pct and d["name"] == client_name
        d_mult = discount_mult if is_discounted else 1.0
        d_pct = discount_pct if is_discounted else 0
        for p in d["products"]:
            info = MARGIN_ASSUMPTIONS[p]
            vol = 1 if p == "Implementation" else d["annual_claim_vol"]
            unit_price = info["unit_price"]
            disc_price = unit_price * d_mult
            ws4[f"A{row_num}"], ws4[f"B{row_num}"], ws4[f"C{row_num}"] = d["name"], p, info["type"]
            ws4[f"D{row_num}"] = unit_price
            ws4[f"E{row_num}"] = f"{d_pct}%" if d_pct else "-"
            ws4[f"F{row_num}"] = round(disc_price, 2)
            ws4[f"G{row_num}"], ws4[f"H{row_num}"] = vol, info["uom"]
            ws4[f"I{row_num}"] = f"=F{row_num}*G{row_num}"
            ws4[f"J{row_num}"], ws4[f"K{row_num}"] = detail_notes.get(p, ""), info["margin"]
            ws4[f"L{row_num}"], ws4[f"M{row_num}"] = f"=I{row_num}*(1-K{row_num})", f"=I{row_num}*K{row_num}"
            row_num += 1

    safe_name = re.sub(r'[^\w\s-]', '', client_name).strip().replace(' ', '_')
    out_path = os.path.join(OUTPUT_DIR, safe_name, f"Sembuh AI - {client_name} - Pricing Internal.xlsx")
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    wb.save(out_path)
    print(f"    Pricing:   {out_path}")
    return out_path


def generate_proposal(deal, discount_pct=0, discount_note="", mom_context=""):
    """Generate proposal using the full dynamic replacement logic from generate_docs."""
    sys.path.insert(0, _BASE_DIR)
    from generate_docs import generate_proposal as _gen_proposal
    out_path = _gen_proposal(deal, discount_pct=discount_pct, discount_note=discount_note, mom_context=mom_context)
    return out_path


# ══════════════════════════════════════════════════════════════
# GOOGLE DRIVE UPLOAD
# ══════════════════════════════════════════════════════════════

def _get_google_tools():
    sys.path.insert(0, _BASE_DIR)
    from google_tools import drive_create_file, drive_upload_file, gmail_create_draft
    return drive_create_file, drive_upload_file, gmail_create_draft


def upload_to_drive(files, client_name):
    drive_create_file, drive_upload_file, _ = _get_google_tools()

    result = drive_create_file(title=client_name, mime_type="folder", folder_id=GOOGLE_FOLDER_ID)
    folder_match = re.search(r"id: ([a-zA-Z0-9_-]+)", result)
    client_folder_id = folder_match.group(1) if folder_match else GOOGLE_FOLDER_ID

    subfolder_map = {"Proposals": "proposal", "Quotations": "quotation", "Pricing": "pricing"}
    for subfolder, keyword in subfolder_map.items():
        sub_result = drive_create_file(title=subfolder, mime_type="folder", folder_id=client_folder_id)
        sub_match = re.search(r"id: ([a-zA-Z0-9_-]+)", sub_result)
        sub_id = sub_match.group(1) if sub_match else client_folder_id
        for f in files:
            if keyword in os.path.basename(f).lower():
                drive_upload_file(f, folder_id=sub_id)
                print(f"    Uploaded {os.path.basename(f)} -> {subfolder}/")

    drive_url = f"https://drive.google.com/drive/folders/{client_folder_id}"
    return client_folder_id, drive_url


# ══════════════════════════════════════════════════════════════
# GMAIL DRAFT
# ══════════════════════════════════════════════════════════════

def create_email_draft(deal, files, discount_pct=0, discount_note=""):
    _, _, gmail_create_draft = _get_google_tools()

    client_name = deal["name"]
    products_str = ", ".join(deal["products"])
    owner_name = deal.get("owner") or "Indi Bintang"
    monthly_vol = int(deal["monthly_claim_vol"] or 0)
    members = int(deal["members_covered"] or 0)

    discount_paragraph = ""
    if discount_pct:
        discount_paragraph = f"""
As discussed, we have applied a {discount_pct}% discount to the proposed commercial terms{(' (' + discount_note + ')') if discount_note else ''}. The updated pricing is reflected in the attached quotation.
"""

    subject = f"Re: Commercial Proposal — {client_name}"
    body = f"""Hi,

Thank you for your time during our recent discussion. We appreciate the opportunity to work with {client_name} and the valuable feedback from your team.
{discount_paragraph}
Following our conversation, please find attached our updated proposal and quotation for your review:

1. Proposal — Covers our recommended solution ({products_str}), implementation approach, and expected outcomes tailored to your operations.
2. Quotation — Detailed pricing breakdown based on your current volume ({monthly_vol:,} claims/month, {members:,} members covered).

Implementation Overview:
- Week 1: Project kickoff, requirements alignment, and environment setup
- Week 2: API integration with your core claims system
- Week 3: End-to-end testing and UAT with your team
- Week 4: Go-live, monitoring, and handover

We're confident that Sembuh AI can deliver meaningful improvements to your claims processing workflow. We'd be happy to walk through the proposal in more detail or adjust any terms to align with your requirements.

Please don't hesitate to reach out if you have any questions or need any clarification. We look forward to your response and the opportunity to move forward together.

Best regards,
{owner_name}
Sembuh AI
www.sembuh.ai
"""

    result = gmail_create_draft(
        to="sales@sembuh.ai", subject=subject, body=body,
        attachments=[f for f in files if "pricing" not in os.path.basename(f).lower()],
    )
    print(f"    Gmail draft: {result}")
    # Extract draft ID for Gmail link
    draft_id_match = re.search(r"ID:\s*(\S+)", result)
    draft_id = draft_id_match.group(1) if draft_id_match else None
    return draft_id


# ══════════════════════════════════════════════════════════════
# SLACK NOTIFICATION (per Slack-Report-Spec.docx)
# ══════════════════════════════════════════════════════════════

def _slack_post(channel, text, blocks):
    """Post message to Slack channel."""
    resp = requests.post(f"{SLACK_API_URL}/chat.postMessage", json={
        "channel": channel, "text": text, "blocks": blocks,
    }, headers={"Authorization": f"Bearer {SLACK_BOT_TOKEN}", "Content-Type": "application/json"})
    data = resp.json()
    if data.get("ok"):
        print(f"  Slack sent to {channel}")
    else:
        print(f"  Slack error ({channel}): {data.get('error')}")
    return data


def _health_emoji(score):
    if score >= 80:
        return ":large_green_circle:"
    if score >= 50:
        return ":warning:"
    if score > 0:
        return ":red_circle:"
    return ":white_circle:"


def _health_label(score):
    if score >= 80:
        return "Healthy"
    if score >= 50:
        return "At Risk"
    if score > 0:
        return "Critical"
    return "Lost"


def slack_managerial_report(deals, generated_files=None, drive_links=None, draft_links=None):
    """Managerial report → #bod-updates (structured card format)."""
    drive_links = drive_links or {}
    draft_links = draft_links or {}

    # Compute health for all deals
    scored = [(d, *compute_health_score(d)) for d in deals]

    active = [d for d in deals if d["stage"] not in ("Lost",)]
    total_value = sum(d["deal_value"] for d in active)
    weighted = sum(d["deal_value"] * STAGE_WEIGHTS.get(d["stage"], 0.1) for d in active)
    prob_pct = (weighted / total_value * 100) if total_value else 0

    at_risk_deals = [(d, s, r) for d, s, st, r in scored if st == "At Risk"]
    critical_deals = [(d, s, r) for d, s, st, r in scored if st == "Critical"]
    healthy_deals = [(d, s) for d, s, st, r in scored if st == "Healthy" and d["stage"] not in ("Won", "Implementation Done", "[Won] Waiting Signature")]

    day_display = TODAY.strftime("%B %d, %Y")

    # ── Block 1: Header ──
    blocks = [
        {"type": "header", "text": {"type": "plain_text", "text": "WEEKLY PIPELINE REPORT"}},
        {"type": "context", "elements": [
            {"type": "mrkdwn", "text": f":calendar:  {day_display}  |  Sembuh AI | Sales Intelligence Engine"},
        ]},
        {"type": "divider"},
    ]

    # ── Block 2: Pipeline Overview ──
    overview = "\n".join([
        "*PIPELINE OVERVIEW*",
        "```",
        f"{'Metric':<22} Value",
        f"{'-'*35}",
        f"{'Active Deals':<22} {len(active)}",
        f"{'Total Pipeline':<22} ${total_value:,.0f}",
        f"{'Weighted Forecast':<22} ${weighted:,.0f} ({prob_pct:.1f}%)",
        f"{'Deals at Risk':<22} {len(at_risk_deals)}",
        f"{'Critical Deals':<22} {len(critical_deals)}",
        "```",
    ])
    blocks.append({"type": "section", "text": {"type": "mrkdwn", "text": overview}})

    # ── Block 3: Deal Summary Table ──
    table_lines = [
        "*DEAL SUMMARY*",
        "```",
        f"{'Deal':<30} {'Stage':<20} {'Value':>10} {'Score':>7} {'Status':<10}",
        f"{'-'*80}",
    ]
    # Sort by score asc (worst first)
    sorted_scored = sorted(scored, key=lambda x: x[1])
    for d, score, status, reasons in sorted_scored:
        if status == "Lost":
            continue
        stage_short = d["stage"][:18]
        st_label = status.upper()
        table_lines.append(f"{d['name'][:29]:<30} {stage_short:<20} ${d['deal_value']:>8,.0f} {score:>5}   {st_label:<10}")
    table_lines.append("```")
    blocks.append({"type": "section", "text": {"type": "mrkdwn", "text": "\n".join(table_lines)}})
    blocks.append({"type": "divider"})

    # ── Block 4: Critical Deals ──
    for d, score, reasons in critical_deals:
        last_activity = _parse_date(d.get("last_contact")) or _parse_date(d.get("updated_at"))
        days_ago = f"{(TODAY - last_activity).days} days ago" if last_activity else "Unknown"
        reason_action = reasons[0] if reasons else "Review immediately"
        card = "\n".join([
            ":red_circle: *CRITICAL DEAL*",
            f"```",
            f"{'Deal':<22} {d['name']}",
            f"{'Value':<22} ${d['deal_value']:,.0f}",
            f"{'Score':<22} {score}/100",
            f"{'Stage':<22} {d['stage']}",
            f"{'Last Activity':<22} {days_ago}",
            f"{'Required Action':<22} {reason_action}",
            f"{'Owner':<22} {d['owner'] or 'Unassigned'}",
            f"```",
        ])
        blocks.append({"type": "section", "text": {"type": "mrkdwn", "text": card}})

    # ── Block 5: At Risk Deals ──
    if at_risk_deals:
        blocks.append({"type": "section", "text": {"type": "mrkdwn", "text": ":warning: *DEALS AT RISK*"}})
        for d, score, reasons in at_risk_deals:
            close = _parse_date(d.get("close_date"))
            close_info = ""
            if close:
                days_left = (close - TODAY).days
                if days_left > 0:
                    close_info = f"{'Close Date':<22} {close.strftime('%B %d')} ({days_left} days remaining)\n"
                else:
                    close_info = f"{'Close Date':<22} {close.strftime('%B %d')} (OVERDUE)\n"

            last_activity = _parse_date(d.get("last_contact")) or _parse_date(d.get("updated_at"))
            stage_info = d["stage"]
            if last_activity:
                days_in = (TODAY - last_activity).days
                if days_in > 0:
                    stage_info = f"{d['stage']} ({days_in} days)"

            gap_line = ""
            for r in reasons:
                if "proposal" in r.lower():
                    gap_line = f"{'Gap':<22} Proposal not submitted\n"
                    break

            reason_action = reasons[0] if reasons else "Follow up this week"
            card = "\n".join(filter(None, [
                f"```",
                f"{'Deal':<22} {d['name']}",
                f"{'Value':<22} ${d['deal_value']:,.0f}",
                f"{'Score':<22} {score}/100",
                f"{close_info.rstrip()}" if close_info else None,
                f"{'Stage':<22} {stage_info}",
                f"{gap_line.rstrip()}" if gap_line else None,
                f"{'Required Action':<22} {reason_action}",
                f"{'Owner':<22} {d['owner'] or 'Unassigned'}",
                f"```",
            ]))
            blocks.append({"type": "section", "text": {"type": "mrkdwn", "text": card}})

    # ── Block 6: Positive Signals ──
    if healthy_deals:
        blocks.append({"type": "divider"})
        blocks.append({"type": "section", "text": {"type": "mrkdwn", "text": ":white_check_mark: *POSITIVE SIGNALS*"}})
        for d, score in healthy_deals:
            card = "\n".join([
                f"```",
                f"{'Deal':<22} {d['name']}",
                f"{'Status':<22} On track",
                f"{'Progress':<22} {d['stage']} — score {score}/100",
                f"```",
            ])
            blocks.append({"type": "section", "text": {"type": "mrkdwn", "text": card}})

    # ── Block 7: AI Insight ──
    # Find stages with multiple stalling deals
    early_stages = ("Open", "First Meeting Done", "Solutioning")
    stalling = [d for d in active if d["stage"] in early_stages]
    stalling_stages = set(d["stage"] for d in stalling)
    if len(stalling) >= 2:
        blocks.append({"type": "divider"})
        stages_str = ", ".join(sorted(stalling_stages))
        insight = "\n".join([
            ":bulb: *AI INSIGHT*",
            f"```",
            f"{'Observation':<22} {len(stalling)} of {len(active)} deals stalling at early stages",
            f"{'Affected Stages':<22} {stages_str}",
            f"{'Recommendation':<22} Implement 7-day follow-up SLA post First Meeting",
            f"{'Objective':<22} Improve conversion and maintain momentum",
            f"```",
        ])
        blocks.append({"type": "section", "text": {"type": "mrkdwn", "text": insight}})

    # ── Block 8: Generated files + links ──
    if generated_files:
        blocks.append({"type": "divider"})
        file_lines = [":page_facing_up: *DOCUMENTS GENERATED*"]
        for deal, files in zip(deals, generated_files):
            if files:
                file_names = ", ".join(f"`{os.path.basename(f)}`" for f in files)
                file_lines.append(f"• {deal['name']}: {file_names}")
                link_parts = []
                if deal["name"] in drive_links:
                    link_parts.append(f"<{drive_links[deal['name']]}|:file_folder: Google Drive>")
                if deal["name"] in draft_links:
                    link_parts.append(f"<{draft_links[deal['name']]}|:envelope: Gmail Draft>")
                if link_parts:
                    file_lines.append(f"   {' · '.join(link_parts)}")
        blocks.append({"type": "section", "text": {"type": "mrkdwn", "text": "\n".join(file_lines)}})

    # ── Footer ──
    blocks.append({"type": "divider"})
    blocks.append({"type": "context", "elements": [
        {"type": "mrkdwn", "text": f"<@{SLACK_NOTIFY_USER}> | _Powered by Sales Intelligence Engine_"},
    ]})

    return _slack_post(SLACK_CHANNEL_BOD, f"<@{SLACK_NOTIFY_USER}> Weekly Pipeline Report — {TODAY_STR}", blocks)


def slack_staff_report(deals, am_name=None, drive_links=None, draft_links=None):
    """Staff report → #am-{name} style (Section 5.1 of spec). Sent to main channel for demo."""
    drive_links = drive_links or {}
    draft_links = draft_links or {}

    if am_name:
        my_deals = [d for d in deals if am_name.lower() in (d.get("owner") or "").lower()]
    else:
        my_deals = deals

    if not my_deals:
        return

    display_name = am_name or "Team"
    day_name = TODAY.strftime("%A, %B %d, %Y")

    lines = [
        f":wave:  Good morning, {display_name}! Here's your deal briefing.",
        f"*{day_name}*",
        "",
        f"*Your Active Deals ({len(my_deals)})*",
        "───────────────────────────",
    ]

    total_pipeline = 0
    urgent_count = 0
    warning_count = 0
    top_priority = None

    for d in my_deals:
        score, status, reasons = compute_health_score(d)
        emoji = _health_emoji(score)
        total_pipeline += d["deal_value"]
        reason_line = "; ".join(reasons) if reasons else "On track"

        if score < 50:
            urgent_count += 1
            if not top_priority:
                top_priority = d["name"]
        elif score < 80:
            warning_count += 1
            if not top_priority:
                top_priority = d["name"]

        lines.append(f"{emoji}  *{d['name']}*")
        lines.append(f"   ${d['deal_value']:,.0f}  |  {d['stage']}  |  Score: {score}/100")

        if status == "Critical":
            lines.append(f"   :rotating_light: {reason_line}")
        elif status == "At Risk":
            lines.append(f"   :warning: {reason_line}")
        else:
            lines.append(f"   :white_check_mark: {reason_line}")

        # Action item
        if score < 50:
            lines.append(f"   :arrow_right: Urgent: Address immediately")
        elif score < 80:
            lines.append(f"   :arrow_right: Follow up this week")
        else:
            lines.append(f"   :arrow_right: Continue current engagement")

        # Links
        link_parts = []
        if d["name"] in drive_links:
            link_parts.append(f"<{drive_links[d['name']]}|:file_folder: Drive>")
        if d["name"] in draft_links:
            link_parts.append(f"<{draft_links[d['name']]}|:envelope: Draft>")
        if link_parts:
            lines.append(f"   {' · '.join(link_parts)}")
        lines.append("")

    lines.append("───────────────────────────")
    lines.append(f"*Summary:* {len(my_deals)} deals | ${total_pipeline:,.0f} pipeline | {urgent_count} urgent, {warning_count} warning")
    if top_priority:
        lines.append(f"*Top priority today:* {top_priority}")

    lines.append("")
    lines.append(f"<@{SLACK_NOTIFY_USER}>")

    text = "\n".join(lines)
    blocks = [{"type": "section", "text": {"type": "mrkdwn", "text": text}}]
    return _slack_post(SLACK_CHANNEL_AM, f"<@{SLACK_NOTIFY_USER}> Daily briefing for {display_name}", blocks)


def slack_critical_alert(deal, score, reasons):
    """Real-time critical alert (Section 4.2 of spec)."""

    reason_str = "; ".join(reasons)
    owner_mention = f"<@{SLACK_NOTIFY_USER}>" if deal.get("owner") else "Unassigned"

    text = "\n".join([
        f":rotating_light:  *Critical Deal Alert*",
        "",
        f"*Deal:* {deal['name']}",
        f"*Value:* ${deal['deal_value']:,.0f}  |  *Stage:* {deal['stage']}",
        f"*Health:* Score dropped to {score}/100 :red_circle:",
        "",
        f"*What happened:* {reason_str}",
        f"*Assigned to:* {owner_mention} ({deal.get('owner') or 'Unassigned'})",
        f"*Recommended:* Review deal and take immediate action.",
        "",
        "_Powered by Sales Intelligence Engine_",
    ])

    blocks = [{"type": "section", "text": {"type": "mrkdwn", "text": text}}]
    return _slack_post(SLACK_CHANNEL_AM, f":rotating_light: Critical: {deal['name']} — {score}/100", blocks)


def slack_deal_activity(deal, event_summary, agent_actions, next_step, drive_url=None, draft_url=None):
    """Real-time deal activity alert (Section 5.2 of spec)."""

    actions_lines = "\n".join(f":white_check_mark: {a}" for a in agent_actions)

    link_lines = []
    if drive_url:
        link_lines.append(f":file_folder: <{drive_url}|Open Google Drive folder>")
    if draft_url:
        link_lines.append(f":envelope: <{draft_url}|Open Gmail draft>")

    text = "\n".join([
        f":incoming_envelope:  *New activity on your deal*",
        "",
        f"*Deal:* {deal['name']} (${deal['deal_value']:,.0f})",
        f"*Event:* {event_summary}",
        "",
        "*Agent actions completed:*",
        actions_lines,
        "",
        *(["*Quick Links:*"] + link_lines + [""] if link_lines else []),
        "*Your next step:*",
        f":arrow_right: {next_step}",
        "",
        f"<@{SLACK_NOTIFY_USER}>",
    ])

    blocks = [{"type": "section", "text": {"type": "mrkdwn", "text": text}}]
    return _slack_post(SLACK_CHANNEL_AM, f":incoming_envelope: Activity: {deal['name']}", blocks)


# ══════════════════════════════════════════════════════════════
# DEAL HEALTH SCORING
# ══════════════════════════════════════════════════════════════

def _parse_date(s):
    """Parse date string (ISO or Monday.com format) to datetime.date, or None."""
    if not s:
        return None
    # Handle ISO datetime with timezone (e.g. "2026-04-24T10:30:00Z")
    s = s.strip()
    for fmt in ("%Y-%m-%d", "%Y-%m-%dT%H:%M:%SZ", "%Y-%m-%dT%H:%M:%S%z"):
        try:
            return datetime.datetime.strptime(s[:25], fmt.replace("%z", "")).date()
        except (ValueError, IndexError):
            continue
    try:
        return datetime.date.fromisoformat(s[:10])
    except ValueError:
        return None


def compute_health_score(deal):
    """Health scoring per Developer Handoff Package Section 4."""
    score = 100
    reasons = []
    stage = deal["stage"]

    if stage in ("Lost",):
        return 0, "Lost", ["Deal marked as Lost."]
    if stage in ("Won", "Implementation Done", "[Won] Waiting Signature"):
        return 100, "Healthy", ["Deal won."]

    # ── Rule 1 & 2: No activity in 7+ / 14+ days (mutually exclusive, take worst) ──
    last_activity = _parse_date(deal.get("last_contact")) or _parse_date(deal.get("updated_at"))
    if last_activity:
        days_inactive = (TODAY - last_activity).days
        if days_inactive >= 14:
            score -= 50
            reasons.append(f"No activity in {days_inactive} days (critical — deal going cold)")
        elif days_inactive >= 7:
            score -= 30
            reasons.append(f"No activity in {days_inactive} days")

    # ── Rule 3: Stuck in same stage 21+ days ──
    # Use updated_at as proxy for stage change date
    stage_date = _parse_date(deal.get("updated_at"))
    if stage_date:
        days_in_stage = (TODAY - stage_date).days
        if days_in_stage >= 21:
            score -= 25
            reasons.append(f"Stuck in {stage} for {days_in_stage} days")

    # ── Rule 4: Expected close date passed ──
    close = _parse_date(deal.get("close_date"))
    if close and close < TODAY:
        score -= 20
        reasons.append(f"Close date {deal['close_date']} has passed")

    # ── Rule 5: No owner assigned ──
    if not deal.get("owner"):
        score -= 15
        reasons.append("No owner assigned — orphaned deal")

    # ── Rule 6: Deal value empty ──
    if deal.get("deal_value", 0) == 0:
        score -= 10
        reasons.append("Deal value empty — can't forecast")

    # ── Rule 7: No proposal sent (Solutioning+ stages) ──
    advanced = ("Solutioning", "Waiting Confirmation", "[Won] Waiting Signature", "Piloting")
    if stage in advanced and not deal.get("proposal_date"):
        score -= 10
        reasons.append(f"Stage is {stage} but no proposal on file")

    # ── Rule 8: High excess rate (>30%) ──
    incurred = deal.get("incurred", 0)
    excess = deal.get("excess", 0)
    if incurred > 0:
        excess_rate = excess / incurred
        if excess_rate > 0.30:
            score -= 10
            reasons.append(f"High excess rate ({excess_rate:.0%}) — claim rejection risk")

    score = max(score, 0)
    if score >= 80:
        status = "Healthy"
    elif score >= 50:
        status = "At Risk"
    else:
        status = "Critical"

    return score, status, reasons


# ══════════════════════════════════════════════════════════════
# SQLite DATABASE
# ══════════════════════════════════════════════════════════════

def _init_db():
    """Initialize SQLite database with deals + action_histories tables."""
    conn = sqlite3.connect(DB_PATH)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS deals (
            id TEXT PRIMARY KEY,
            name TEXT NOT NULL,
            stage TEXT,
            deal_value REAL DEFAULT 0,
            close_date TEXT,
            monthly_claim_vol REAL DEFAULT 0,
            annual_claim_vol REAL DEFAULT 0,
            members_covered REAL DEFAULT 0,
            pricing_model TEXT,
            incurred REAL DEFAULT 0,
            excess REAL DEFAULT 0,
            approved REAL DEFAULT 0,
            op_pct REAL DEFAULT 0,
            ip_pct REAL DEFAULT 0,
            owner TEXT,
            proposal_date TEXT,
            products TEXT,
            health_score INTEGER DEFAULT 0,
            health_status TEXT,
            health_reasons TEXT,
            synced_at TEXT
        )
    """)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS action_histories (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            deal_id TEXT,
            deal_name TEXT NOT NULL,
            action_type TEXT NOT NULL,
            description TEXT,
            draft_link TEXT,
            drive_link TEXT,
            files_generated TEXT,
            mom_file TEXT,
            discount_pct REAL DEFAULT 0,
            status TEXT DEFAULT 'pending',
            created_at TEXT NOT NULL,
            processed_at TEXT
        )
    """)
    conn.commit()
    return conn


def _sync_deals_to_db():
    """Fetch deals from Monday.com, compute health, store in SQLite. Returns deal list."""
    boards = fetch_boards()
    deals_board = find_deals_board(boards)
    if not deals_board:
        raise RuntimeError("No Deals board found")

    raw_deals = fetch_deals(deals_board["id"])
    deals = [parse_deal(d) for d in raw_deals]
    now = datetime.datetime.now().isoformat()

    conn = _init_db()
    for deal in deals:
        score, status, reasons = compute_health_score(deal)
        deal["health_score"] = score
        deal["health_status"] = status
        deal["health_reasons"] = reasons

        conn.execute("""
            INSERT OR REPLACE INTO deals
            (id, name, stage, deal_value, close_date, monthly_claim_vol, annual_claim_vol,
             members_covered, pricing_model, incurred, excess, approved, op_pct, ip_pct,
             owner, proposal_date, products, health_score, health_status, health_reasons, synced_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            deal["id"], deal["name"], deal["stage"], deal["deal_value"],
            deal["close_date"], deal["monthly_claim_vol"], deal["annual_claim_vol"],
            deal["members_covered"], deal["pricing_model"], deal["incurred"],
            deal["excess"], deal["approved"], deal.get("op_pct", 0), deal.get("ip_pct", 0),
            deal["owner"], deal.get("proposal_date", ""),
            json.dumps(deal["products"]),
            score, status, json.dumps(reasons), now,
        ))
    conn.commit()
    conn.close()
    return deals, now


def _get_deals_from_db():
    """Read all deals from SQLite, return list of dicts."""
    conn = _init_db()
    conn.row_factory = sqlite3.Row
    rows = conn.execute("SELECT * FROM deals ORDER BY health_score ASC").fetchall()
    conn.close()
    deals = []
    for row in rows:
        d = dict(row)
        d["products"] = json.loads(d["products"]) if d["products"] else []
        d["health_reasons"] = json.loads(d["health_reasons"]) if d["health_reasons"] else []
        deals.append(d)
    return deals


def _log_action(deal, action_type, description, draft_link="", drive_link="",
                files_generated=None, mom_file="", discount_pct=0):
    """Insert action history record."""
    conn = _init_db()
    now = datetime.datetime.now().isoformat()
    conn.execute("""
        INSERT INTO action_histories
        (deal_id, deal_name, action_type, description, draft_link, drive_link,
         files_generated, mom_file, discount_pct, status, created_at)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, 'pending', ?)
    """, (
        deal.get("id", ""), deal["name"], action_type, description,
        draft_link, drive_link,
        json.dumps(files_generated or []), mom_file, discount_pct, now,
    ))
    conn.commit()
    conn.close()


def _get_pending_actions():
    """Get all unprocessed action histories."""
    conn = _init_db()
    conn.row_factory = sqlite3.Row
    rows = conn.execute(
        "SELECT * FROM action_histories WHERE status = 'pending' ORDER BY created_at DESC"
    ).fetchall()
    conn.close()
    actions = []
    for row in rows:
        d = dict(row)
        d["files_generated"] = json.loads(d["files_generated"]) if d["files_generated"] else []
        actions.append(d)
    return actions


def _get_all_actions(limit=50):
    """Get recent action histories (all statuses)."""
    conn = _init_db()
    conn.row_factory = sqlite3.Row
    rows = conn.execute(
        "SELECT * FROM action_histories ORDER BY created_at DESC LIMIT ?", (limit,)
    ).fetchall()
    conn.close()
    actions = []
    for row in rows:
        d = dict(row)
        d["files_generated"] = json.loads(d["files_generated"]) if d["files_generated"] else []
        actions.append(d)
    return actions


def _mark_action_done(action_id):
    """Mark action history as processed."""
    conn = _init_db()
    now = datetime.datetime.now().isoformat()
    conn.execute(
        "UPDATE action_histories SET status = 'done', processed_at = ? WHERE id = ?",
        (now, action_id)
    )
    conn.commit()
    conn.close()


def _build_dashboard_data(deals):
    """Build full dashboard payload from deal list."""
    active = [d for d in deals if d["stage"] not in ("Lost",)]
    lost = [d for d in deals if d["stage"] == "Lost"]

    total_value = sum(d["deal_value"] for d in active)
    weighted = sum(d["deal_value"] * STAGE_WEIGHTS.get(d["stage"], 0.1) for d in active)
    prob_pct = int((weighted / total_value * 100) if total_value else 0)

    healthy = [d for d in deals if d["health_status"] == "Healthy"]
    at_risk = [d for d in deals if d["health_status"] == "At Risk"]
    critical = [d for d in deals if d["health_status"] == "Critical"]

    # Find worst critical deal for alert subtitle
    critical_subtitle = ""
    if critical:
        worst = min(critical, key=lambda d: d["health_score"])
        reasons = worst["health_reasons"]
        critical_subtitle = f"{worst['name']} — {reasons[0]}" if reasons else worst["name"]

    # KPI
    kpi = {
        "total_pipeline": total_value,
        "active_deals": len(active),
        "weighted_pipeline": weighted,
        "conversion_probability": prob_pct,
        "at_risk_count": len(at_risk),
        "critical_count": len(critical),
        "critical_subtitle": critical_subtitle,
    }

    # Deal health table
    deal_table = []
    for d in deals:
        deal_table.append({
            "name": d["name"],
            "stage": d["stage"],
            "value": d["deal_value"],
            "score": d["health_score"],
            "status": d["health_status"],
            "reasons": d["health_reasons"],
            "owner": d["owner"],
            "close_date": d["close_date"],
        })

    # Pipeline by stage
    stage_groups = {}
    for d in deals:
        s = d["stage"]
        if s not in stage_groups:
            stage_groups[s] = {"value": 0, "count": 0}
        stage_groups[s]["value"] += d["deal_value"]
        stage_groups[s]["count"] += 1

    max_val = max((v["value"] for v in stage_groups.values()), default=1)
    pipeline_stages = []
    for stage, info in sorted(stage_groups.items(), key=lambda x: x[1]["value"]):
        pipeline_stages.append({
            "stage": stage,
            "value": info["value"],
            "count": info["count"],
            "pct": round(info["value"] / max_val * 100) if max_val else 0,
        })

    # AI actions
    actions = []
    scored = [(d, d["health_score"], d["health_status"], d["health_reasons"]) for d in deals]
    scored.sort(key=lambda x: x[1])
    for d, score, status, reasons in scored:
        if status == "Lost":
            continue
        if status == "Critical":
            actions.append({
                "priority": "urgent",
                "text": reasons[0] if reasons else f"{d['name']} needs immediate attention",
                "deal": f"{d['name']} • ${d['deal_value']:,.0f}",
            })
        elif status == "At Risk":
            actions.append({
                "priority": "warning",
                "text": reasons[0] if reasons else f"{d['name']} at risk",
                "deal": f"{d['name']} • ${d['deal_value']:,.0f}",
            })
        else:
            actions.append({
                "priority": "info",
                "text": "On track — continue current engagement",
                "deal": f"{d['name']} • ${d['deal_value']:,.0f}",
            })

    # Revenue forecast (bar chart data)
    revenue_forecast = {
        "labels": [],
        "deal_values": [],
        "weighted_values": [],
    }
    for d in active:
        short = d["name"][:20]
        weight = STAGE_WEIGHTS.get(d["stage"], 0.1)
        revenue_forecast["labels"].append(short)
        revenue_forecast["deal_values"].append(d["deal_value"])
        revenue_forecast["weighted_values"].append(round(d["deal_value"] * weight))

    # Health distribution (doughnut)
    health_distribution = {
        "labels": ["Healthy (80-100)", "At Risk (50-79)", "Critical (<50)", "Lost"],
        "values": [len(healthy), len(at_risk), len(critical), len(lost)],
    }

    # AI narrative
    narrative_parts = [
        f"Total active pipeline stands at ${total_value:,.0f} across {len(active)} deals"
        f" (excluding {len(lost)} lost).",
        f"Weighted forecast is ${weighted:,.0f} based on stage probabilities ({prob_pct}% conversion).",
    ]
    if critical:
        names = ", ".join(d["name"] for d in critical)
        narrative_parts.append(f"Critical: {names} — require immediate action.")
    if at_risk:
        names = ", ".join(d["name"] for d in at_risk)
        narrative_parts.append(f"At risk: {names} — monitor closely this week.")
    if healthy:
        names = ", ".join(d["name"] for d in healthy)
        narrative_parts.append(f"Positive: {names} — progressing well.")

    synced_at = deals[0].get("synced_at", "") if deals else ""

    return {
        "kpi": kpi,
        "deals": deal_table,
        "pipeline_stages": pipeline_stages,
        "actions": actions,
        "revenue_forecast": revenue_forecast,
        "health_distribution": health_distribution,
        "narrative": " ".join(narrative_parts),
        "synced_at": synced_at,
    }


# ══════════════════════════════════════════════════════════════
# FLASK API SERVER
# ══════════════════════════════════════════════════════════════

def create_app():
    """Create Flask app with API endpoints."""
    from flask import Flask, jsonify, request, send_from_directory
    from flask_cors import CORS

    app = Flask(__name__, static_folder=os.path.join(_BASE_DIR, "PTP Hackathon - Brief"))
    CORS(app)

    @app.route("/api/sync", methods=["POST"])
    def api_sync():
        """Sync deals from Monday.com → SQLite."""
        try:
            deals, synced_at = _sync_deals_to_db()
            return jsonify({
                "status": "ok",
                "deals_synced": len(deals),
                "synced_at": synced_at,
            })
        except Exception as e:
            return jsonify({"status": "error", "message": str(e)}), 500

    @app.route("/api/dashboard", methods=["GET"])
    def api_dashboard():
        """Get dashboard data from SQLite."""
        deals = _get_deals_from_db()
        if not deals:
            return jsonify({"status": "empty", "message": "No data. Call POST /api/sync first."}), 404
        data = _build_dashboard_data(deals)
        return jsonify(data)

    @app.route("/api/actions", methods=["GET"])
    def api_actions():
        """Get action histories. ?status=pending for unprocessed only."""
        status_filter = request.args.get("status", "")
        if status_filter == "pending":
            actions = _get_pending_actions()
        else:
            actions = _get_all_actions()
        return jsonify({"actions": actions, "total": len(actions)})

    @app.route("/api/actions/<int:action_id>/done", methods=["POST"])
    def api_action_done(action_id):
        """Mark action as processed."""
        _mark_action_done(action_id)
        return jsonify({"status": "ok", "id": action_id})

    @app.route("/api/pipeline", methods=["GET"])
    def api_pipeline():
        """Run pipeline for a deal. Params: ?deal=Sentosa&mom_path=optional"""
        deal_filter = request.args.get("deal", "")
        mom_path = request.args.get("mom_path", "")

        try:
            result = _run_pipeline_core(deal_filter=deal_filter, docx_path=mom_path)
            return jsonify(result)
        except Exception as e:
            return jsonify({"status": "error", "message": str(e)}), 500

    @app.route("/")
    def index():
        return send_from_directory(_BASE_DIR, "dashboard.html")

    return app


def _run_pipeline_core(deal_filter="", docx_path=""):
    """Core pipeline logic: fetch deals → generate docs → Drive → Gmail → Slack → log actions.
    Returns result dict. Used by both CLI cmd_pipeline and API /api/pipeline.
    """
    result = {"status": "ok", "steps": [], "deals_processed": 0, "files_generated": 0}

    # ── STEP 1: MoM extraction (optional) ──
    mom_text = None
    discount_pct = 0
    discount_note = ""
    mom_context = ""

    if docx_path and os.path.exists(docx_path):
        print(f"  Extracting MoM: {docx_path}")
        mom_text = extract_docx_text(docx_path)
        result["steps"].append(f"MoM extracted ({len(mom_text)} chars)")

        # Extract metadata
        meta = extract_mom_metadata(mom_text)
        discount_pct = meta.get("discount_pct") or 0
        discount_note = meta.get("discount_note") or ""
        mom_context = meta.get("mom_context") or ""
        if discount_pct:
            result["steps"].append(f"Discount detected: {discount_pct}%")

    # ── STEP 2: Fetch deals ──
    boards = fetch_boards()
    deals_board = find_deals_board(boards)
    if not deals_board:
        raise RuntimeError("No Deals board found")

    raw_deals = fetch_deals(deals_board["id"])
    all_deals = [parse_deal(d) for d in raw_deals]

    deals = all_deals
    if deal_filter:
        deals = [d for d in all_deals if deal_filter.lower() in d["name"].lower()]
        if not deals:
            raise RuntimeError(f"No deal matching '{deal_filter}'")

    result["deals_processed"] = len(deals)
    result["steps"].append(f"Fetched {len(deals)} deal(s) from Monday.com")

    # ── STEP 3: Generate docs ──
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    all_files = []
    for i, deal in enumerate(deals):
        print(f"  [{deal['name']}] generating docs...")
        files = [
            generate_quotation(deal, seq_num=i + 1, discount_pct=discount_pct, discount_note=discount_note),
            generate_pricing_internal(deal, deals, discount_pct=discount_pct, discount_note=discount_note),
            generate_proposal(deal, discount_pct=discount_pct, discount_note=discount_note, mom_context=mom_context),
        ]
        all_files.append(files)

    total_files = sum(len(f) for f in all_files)
    result["files_generated"] = total_files
    result["steps"].append(f"Generated {total_files} documents")

    # ── STEP 4: Upload to Google Drive ──
    drive_links = {}
    for deal, files in zip(deals, all_files):
        try:
            _, drive_url = upload_to_drive(files, deal["name"])
            drive_links[deal["name"]] = drive_url
        except Exception as e:
            print(f"    Drive error ({deal['name']}): {e}")
    if drive_links:
        result["steps"].append(f"Uploaded to Google Drive ({len(drive_links)} deals)")

    # ── STEP 5: Gmail drafts ──
    draft_links = {}
    for deal, files in zip(deals, all_files):
        if deal["stage"] == "Lost":
            continue
        try:
            draft_id = create_email_draft(deal, files, discount_pct=discount_pct, discount_note=discount_note)
            if draft_id:
                draft_links[deal["name"]] = f"https://mail.google.com/mail/u/0/#drafts?compose={draft_id}"
        except Exception as e:
            print(f"    Gmail error ({deal['name']}): {e}")
    if draft_links:
        result["steps"].append(f"Gmail drafts created ({len(draft_links)} deals)")

    # ── STEP 6: Log action histories ──
    for deal, files in zip(deals, all_files):
        file_basenames = [os.path.basename(f) for f in files]
        draft_link = draft_links.get(deal["name"], "")
        drive_link = drive_links.get(deal["name"], "")
        desc_parts = []
        if files:
            desc_parts.append(f"Generated {len(files)} documents")
        if drive_link:
            desc_parts.append("Uploaded to Google Drive")
        if draft_link:
            desc_parts.append("Gmail draft created")
        if discount_pct:
            desc_parts.append(f"Discount {discount_pct}% applied")
        description = ". ".join(desc_parts) + "." if desc_parts else "Pipeline executed."

        _log_action(
            deal=deal,
            action_type="pipeline",
            description=description,
            draft_link=draft_link,
            drive_link=drive_link,
            files_generated=file_basenames,
            mom_file=os.path.basename(docx_path) if mom_text else "",
            discount_pct=discount_pct,
        )
    result["steps"].append(f"Logged {len(deals)} action(s)")

    # ── STEP 7: Sync deals to DB ──
    _sync_deals_to_db()
    result["steps"].append("Synced deals to dashboard DB")

    # ── STEP 8: Slack notifications ──
    try:
        slack_managerial_report(all_deals, all_files, drive_links=drive_links, draft_links=draft_links)
        owners = set(d["owner"] for d in all_deals if d["owner"])
        for owner in owners:
            slack_staff_report(all_deals, am_name=owner, drive_links=drive_links, draft_links=draft_links)
        if not owners:
            slack_staff_report(all_deals, drive_links=drive_links, draft_links=draft_links)
        for deal in all_deals:
            score, status, reasons = compute_health_score(deal)
            if score < 50 and status == "Critical":
                slack_critical_alert(deal, score, reasons)
        result["steps"].append("Slack notifications sent")
    except Exception as e:
        result["steps"].append(f"Slack error: {e}")

    # ── Result summary ──
    result["drive_links"] = drive_links
    result["draft_links"] = draft_links
    result["deal_names"] = [d["name"] for d in deals]

    return result


# ══════════════════════════════════════════════════════════════
# MoM EXTRACTION
# ══════════════════════════════════════════════════════════════

def extract_docx_text(filepath):
    from docx import Document
    doc = Document(filepath)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for i, table in enumerate(doc.tables):
        parts.append(f"\n=== TABLE {i + 1} ===")
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts)


def extract_mom_metadata(mom_text):
    """Extract discount, context summary, and other commercial metadata from MoM using Claude."""
    anthropic = Anthropic()
    resp = anthropic.messages.create(
        model=CLAUDE_MODEL,
        max_tokens=512,
        messages=[{"role": "user", "content": f"""\
Extract commercial metadata from this meeting notes / email. Return ONLY valid JSON:

--- TEXT ---
{mom_text}
--- END ---

{{
  "discount_pct": <number 0-100 or null if no discount mentioned>,
  "discount_note": "<reason for discount or null>",
  "mom_context": "<1-2 sentence summary of key outcomes, decisions, and next steps from this meeting — suitable for including in a proposal document>",
  "client_contact_name": "<name of the main client contact/stakeholder mentioned or null>",
  "client_contact_title": "<job title of the client contact or null>",
  "client_company": "<client company name or null>"
}}
"""}],
    )
    raw = resp.content[0].text.strip()
    if raw.startswith("```"):
        raw = raw.split("\n", 1)[1].rsplit("```", 1)[0].strip()
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        return {"discount_pct": None, "discount_note": None, "mom_context": None, "client_contact_name": None, "client_contact_title": None, "client_company": None}


# ══════════════════════════════════════════════════════════════
# CLI COMMANDS
# ══════════════════════════════════════════════════════════════

def cmd_generate(args):
    """Fetch deals from Monday.com, generate all documents, upload, email, notify."""
    print("=" * 60)
    print("Sales Intelligence — Document Generator")
    print("=" * 60)

    # 1. Fetch deals
    print("\n[1/5] Fetching deals from Monday.com...")
    boards = fetch_boards()
    deals_board = find_deals_board(boards)
    if not deals_board:
        print("No board found."); sys.exit(1)

    print(f"  Board: {deals_board['name']} (id: {deals_board['id']})")
    raw_deals = fetch_deals(deals_board["id"])
    all_deals = [parse_deal(d) for d in raw_deals]

    # Filter by --deal for doc generation, keep all_deals for reporting
    deals = all_deals
    if args.deal:
        deals = [d for d in all_deals if args.deal.lower() in d["name"].lower()]
        if not deals:
            print(f"  No deal matching '{args.deal}'"); sys.exit(1)

    print(f"  Processing {len(deals)} deals (total in CRM: {len(all_deals)})")

    # 2. Generate docs
    print("\n[2/5] Generating documents...")
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    all_files = []
    for i, deal in enumerate(deals):
        print(f"\n  [{deal['name']}] ({deal['stage']})")
        files = [
            generate_quotation(deal, seq_num=i + 1),
            generate_pricing_internal(deal, deals),
            generate_proposal(deal),
        ]
        all_files.append(files)

    # 3. Upload to Google Drive
    print("\n[3/5] Uploading to Google Drive...")
    drive_links = {}
    for deal, files in zip(deals, all_files):
        try:
            _, drive_url = upload_to_drive(files, deal["name"])
            drive_links[deal["name"]] = drive_url
        except Exception as e:
            print(f"    Drive error ({deal['name']}): {e}")

    # 4. Gmail drafts
    print("\n[4/5] Creating Gmail drafts...")
    draft_links = {}
    for deal, files in zip(deals, all_files):
        if deal["stage"] == "Lost":
            print(f"    Skip {deal['name']} (Lost)"); continue
        try:
            draft_id = create_email_draft(deal, files)
            if draft_id:
                draft_links[deal["name"]] = f"https://mail.google.com/mail/u/0/#drafts?compose={draft_id}"
        except Exception as e:
            print(f"    Gmail error ({deal['name']}): {e}")

    # 5. Slack — managerial uses ALL deals, staff/critical use filtered
    print("\n[5/5] Sending Slack reports...")
    try:
        slack_managerial_report(all_deals, all_files, drive_links=drive_links, draft_links=draft_links)
        owners = set(d["owner"] for d in all_deals if d["owner"])
        for owner in owners:
            slack_staff_report(all_deals, am_name=owner, drive_links=drive_links, draft_links=draft_links)
        if not owners:
            slack_staff_report(all_deals, drive_links=drive_links, draft_links=draft_links)
        for deal in all_deals:
            score, status, reasons = compute_health_score(deal)
            if score < 50 and status == "Critical":
                slack_critical_alert(deal, score, reasons)
    except Exception as e:
        print(f"    Slack error: {e}")

    print(f"\nDone. Output: {OUTPUT_DIR}")


def cmd_mom(args):
    """Extract MoM from docx and update Monday.com CRM via MCP."""
    if not os.path.exists(args.docx_path):
        print(f"File not found: {args.docx_path}"); sys.exit(1)

    print(f"Extracting MoM: {args.docx_path}")
    mom_text = extract_docx_text(args.docx_path)
    print(f"  {len(mom_text)} chars extracted")

    if args.dry_run:
        print("\n=== DRY RUN ===\n")
        print(mom_text)
        return

    async def _run():
        client = MondayMCPClient()
        print("Connecting to Monday MCP...", file=sys.stderr)
        await client.connect()

        system = (
            f"You are a Sales Intelligence CRM updater. "
            f"Workspace ID: {MONDAY_WORKSPACE_ID}. "
            f"Always pass workspace_ids=[{MONDAY_WORKSPACE_ID}] when listing boards. "
            "Analyze MoM data, find matching deal, update CRM fields, add comment summary."
        )

        prompt = f"""\
Here is the extracted Minutes of Meeting (MoM). Update Monday.com CRM accordingly.

--- MoM ---
{mom_text}
--- End ---

Board & Column Reference (Deals board ID from workspace):
  - name: Deal name
  - deal_stage: Status (Waiting Confirmation, Won, Lost, Solutioning, First Meeting Done, Piloting)
  - deal_value: Deal Value (numbers)
  - deal_expected_close_date: Close Date (date format: {{"date":"YYYY-MM-DD"}})
  - deal_close_probability: Close Probability (numbers)

Instructions:
1. List boards in workspace to find Deals board.
2. Search for the deal matching this MoM client.
3. Update relevant columns from MoM data.
4. Add a comment summarizing key MoM points (meeting date, attendees, summary, action items, next steps).
   Keep comment under 2000 chars, plain ASCII only.
"""

        try:
            result, _ = await client.chat(prompt, [], system=system)
            print(f"\n{result}")
        finally:
            await client.cleanup()

    asyncio.run(_run())


def cmd_health(args):
    """Run deal health scoring on all deals."""
    print("Sales Intelligence — Deal Health Monitor\n")

    boards = fetch_boards()
    deals_board = find_deals_board(boards)
    raw_deals = fetch_deals(deals_board["id"])
    deals = [parse_deal(d) for d in raw_deals]

    print(f"{'Deal':<40} {'Score':>5} {'Status':<10} Reasons")
    print("-" * 100)

    critical_deals = []
    for deal in deals:
        score, status, reasons = compute_health_score(deal)
        reason_str = "; ".join(reasons) if reasons else "On track"
        indicator = {"Healthy": "G", "At Risk": "A", "Critical": "R", "Lost": "X"}.get(status, "?")
        print(f"[{indicator}] {deal['name']:<37} {score:>5} {status:<10} {reason_str}")
        if score < 50:
            critical_deals.append((deal, score, reasons))

    # Slack alerts for critical deals
    if critical_deals and SLACK_BOT_TOKEN:
        print(f"\nSending Slack alerts for {len(critical_deals)} critical deals...")
        for deal, score, reasons in critical_deals:
            slack_critical_alert(deal, score, reasons)

    # Full managerial report if requested
    if SLACK_BOT_TOKEN:
        print("\nSending managerial health report...")
        slack_managerial_report(deals)


def cmd_pipeline(args):
    """Full pipeline: MoM → CRM update → generate docs → Drive → Gmail → Slack. One command."""

    print("=" * 60)
    print("Sales Intelligence — Full Pipeline")
    print("=" * 60)

    docx_path = args.docx_path

    # ── STEP 1: Extract MoM & update Monday.com CRM ──────────
    print(f"\n{'='*60}")
    print("STEP 1/3 — MoM Extraction & CRM Update")
    print("=" * 60)

    if not os.path.exists(docx_path):
        print(f"  MoM file not found: {docx_path}")
        print("  Skipping MoM step. Proceeding with existing CRM data.\n")
        mom_text = None
    else:
        print(f"  Extracting: {docx_path}")
        mom_text = extract_docx_text(docx_path)
        print(f"  {len(mom_text)} chars extracted")

        async def _update_crm():
            client = MondayMCPClient()
            print("  Connecting to Monday MCP...", file=sys.stderr)
            await client.connect()

            system = (
                f"You are a Sales Intelligence CRM updater. "
                f"Workspace ID: {MONDAY_WORKSPACE_ID}. "
                f"Always pass workspace_ids=[{MONDAY_WORKSPACE_ID}] when listing boards. "
                "Analyze MoM data, find matching deal, update CRM fields, add comment summary."
            )

            prompt = f"""\
Here is the extracted Minutes of Meeting (MoM). Update Monday.com CRM accordingly.

--- MoM ---
{mom_text}
--- End ---

Board & Column Reference (Deals board ID from workspace):
  - name: Deal name
  - deal_stage: Status (Waiting Confirmation, Won, Lost, Solutioning, First Meeting Done, Piloting)
  - deal_value: Deal Value (numbers)
  - deal_expected_close_date: Close Date (date format: {{"date":"YYYY-MM-DD"}})
  - deal_close_probability: Close Probability (numbers)

Instructions:
1. List boards in workspace to find Deals board.
2. Search for the deal matching this MoM client.
3. Update relevant columns from MoM data.
4. Add a comment summarizing key MoM points (meeting date, attendees, summary, action items, next steps).
   Keep comment under 2000 chars, plain ASCII only.
"""
            try:
                result, _ = await client.chat(prompt, [], system=system)
                print(f"\n  CRM Update Result:\n{result}")
            finally:
                await client.cleanup()

        try:
            asyncio.run(_update_crm())
            print("\n  CRM update complete.")
        except Exception as e:
            print(f"\n  CRM update error: {e}")
            print("  Continuing with document generation...\n")

    # ── Extract discount + context + client contact from MoM ──
    discount_pct = 0
    discount_note = ""
    mom_context = ""
    client_contact_name = ""
    client_contact_title = ""
    if mom_text:
        print("\n  Extracting commercial metadata from MoM...")
        meta = extract_mom_metadata(mom_text)
        discount_pct = meta.get("discount_pct") or 0
        discount_note = meta.get("discount_note") or ""
        mom_context = meta.get("mom_context") or ""
        client_contact_name = meta.get("client_contact_name") or ""
        client_contact_title = meta.get("client_contact_title") or ""
        if discount_pct:
            print(f"  Discount detected: {discount_pct}% — {discount_note}")
        else:
            print("  No discount mentioned in MoM.")
        if client_contact_name:
            print(f"  Client contact: {client_contact_name}" + (f" ({client_contact_title})" if client_contact_title else ""))
        if mom_context:
            print(f"  MoM context: {mom_context}")

    # ── STEP 2: Generate Documents (Point 2 & 3) ─────────────
    print(f"\n{'='*60}")
    print("STEP 2/3 — Document Generation")
    print("=" * 60)

    print("\n  Fetching deals from Monday.com...")
    boards = fetch_boards()
    deals_board = find_deals_board(boards)
    if not deals_board:
        print("  No board found."); sys.exit(1)

    print(f"  Board: {deals_board['name']} (id: {deals_board['id']})")
    raw_deals = fetch_deals(deals_board["id"])
    all_deals = [parse_deal(d) for d in raw_deals]

    # Filter for doc generation, keep all_deals for reporting
    deals = all_deals
    if args.deal:
        deals = [d for d in all_deals if args.deal.lower() in d["name"].lower()]
        if not deals:
            print(f"  No deal matching '{args.deal}'"); sys.exit(1)

    print(f"  Generating docs for {len(deals)} deal(s) (total in CRM: {len(all_deals)})\n")

    os.makedirs(OUTPUT_DIR, exist_ok=True)
    all_files = []
    for i, deal in enumerate(deals):
        print(f"  [{deal['name']}] ({deal['stage']})")
        files = [
            generate_quotation(deal, seq_num=i + 1, discount_pct=discount_pct, discount_note=discount_note, client_contact_name=client_contact_name, client_contact_title=client_contact_title, mom_context=mom_context),
            generate_pricing_internal(deal, deals, discount_pct=discount_pct, discount_note=discount_note),
            generate_proposal(deal, discount_pct=discount_pct, discount_note=discount_note, mom_context=mom_context),
        ]
        all_files.append(files)

    # Upload to Google Drive
    print("\n  Uploading to Google Drive...")
    drive_links = {}
    for deal, files in zip(deals, all_files):
        try:
            _, drive_url = upload_to_drive(files, deal["name"])
            drive_links[deal["name"]] = drive_url
        except Exception as e:
            print(f"    Drive error ({deal['name']}): {e}")

    # Gmail drafts
    print("\n  Creating Gmail drafts...")
    draft_links = {}
    for deal, files in zip(deals, all_files):
        if deal["stage"] == "Lost":
            print(f"    Skip {deal['name']} (Lost)"); continue
        try:
            draft_id = create_email_draft(deal, files, discount_pct=discount_pct, discount_note=discount_note)
            if draft_id:
                draft_links[deal["name"]] = f"https://mail.google.com/mail/u/0/#drafts?compose={draft_id}"
        except Exception as e:
            print(f"    Gmail error ({deal['name']}): {e}")

    # ── Log action histories ──
    print("\n  Logging action histories...")
    for deal, files in zip(deals, all_files):
        file_basenames = [os.path.basename(f) for f in files]
        draft_link = draft_links.get(deal["name"], "")
        drive_link = drive_links.get(deal["name"], "")
        desc_parts = []
        if files:
            desc_parts.append(f"Generated {len(files)} documents")
        if drive_link:
            desc_parts.append("Uploaded to Google Drive")
        if draft_link:
            desc_parts.append("Gmail draft created")
        if discount_pct:
            desc_parts.append(f"Discount {discount_pct}% applied")
        description = ". ".join(desc_parts) + "." if desc_parts else "Pipeline executed."

        _log_action(
            deal=deal,
            action_type="pipeline",
            description=description,
            draft_link=draft_link,
            drive_link=drive_link,
            files_generated=file_basenames,
            mom_file=os.path.basename(docx_path) if mom_text else "",
            discount_pct=discount_pct,
        )
    print(f"  {len(deals)} action(s) logged.")

    # ── STEP 3: Health Check & Slack Notification ─────────────
    print(f"\n{'='*60}")
    print("STEP 3/3 — Health Check & Notifications")
    print("=" * 60)

    # Health scores for ALL deals
    print("\n  Deal Health Scores:")
    print(f"  {'Deal':<35} {'Score':>5} {'Status':<10} Reasons")
    print(f"  {'-'*85}")
    critical_deals = []
    for deal in all_deals:
        score, status, reasons = compute_health_score(deal)
        reason_str = "; ".join(reasons) if reasons else "On track"
        ind = {"Healthy": "G", "At Risk": "A", "Critical": "R", "Lost": "X"}.get(status, "?")
        print(f"  [{ind}] {deal['name']:<32} {score:>5} {status:<10} {reason_str}")
        if score < 50:
            critical_deals.append((deal, score, reasons))

    # Slack reports — managerial/staff always use ALL deals
    print("\n  Sending Slack reports...")

    # 1. Managerial report — all deals
    slack_managerial_report(all_deals, all_files, drive_links=drive_links, draft_links=draft_links)

    # 2. Staff reports per AM — all deals
    owners = set(d["owner"] for d in all_deals if d["owner"])
    for owner in owners:
        slack_staff_report(all_deals, am_name=owner, drive_links=drive_links, draft_links=draft_links)
    if not owners:
        slack_staff_report(all_deals, drive_links=drive_links, draft_links=draft_links)

    # 3. Critical alerts — all deals
    for deal, score, reasons in critical_deals:
        slack_critical_alert(deal, score, reasons)

    # 4. Deal activity alert (for MoM-processed deal)
    if mom_text:
        mom_deal = deals[0] if len(deals) == 1 else None
        if mom_deal:
            deal_drive = drive_links.get(mom_deal["name"])
            deal_draft = draft_links.get(mom_deal["name"])
            slack_deal_activity(
                mom_deal,
                event_summary=f"MoM processed from `{os.path.basename(docx_path)}`",
                agent_actions=[
                    "MoM extracted and parsed",
                    "CRM updated with meeting notes",
                    f"Documents generated (Quotation + Proposal + Pricing)",
                    "Gmail draft created",
                    "Files uploaded to Google Drive",
                ],
                next_step="Review Gmail draft and send to client",
                drive_url=deal_drive,
                draft_url=deal_draft,
            )

    # Summary
    print(f"\n{'='*60}")
    print("PIPELINE COMPLETE")
    print(f"  MoM:       {'Processed' if mom_text else 'Skipped'}")
    print(f"  Deals:     {len(deals)}")
    print(f"  Files:     {sum(len(f) for f in all_files)}")
    print(f"  Output:    {OUTPUT_DIR}")
    print(f"  Slack:     {SLACK_CHANNEL_BOD} (#bod-updates), {SLACK_CHANNEL_AM} (#am-indi)")
    print("=" * 60)


def cmd_interactive(args):
    """Interactive Monday.com assistant."""
    async def _run():
        client = MondayMCPClient()
        print("Monday.com AI Assistant")
        print("Connecting...", file=sys.stderr)
        await client.connect()
        print("Ready. Type 'quit' to exit.\n")

        history = []
        while True:
            try:
                prompt = input("> ").strip()
            except (EOFError, KeyboardInterrupt):
                print("\nBye!"); break
            if not prompt:
                continue
            if prompt.lower() in ("quit", "exit", "q"):
                print("Bye!"); break

            try:
                text, history = await client.chat(prompt, history)
                print(f"\n{text}\n")
            except Exception as e:
                print(f"\nError: {e}\n", file=sys.stderr)

        await client.cleanup()

    asyncio.run(_run())


# ══════════════════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(
        description="Sales Intelligence — Unified Automation",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""\
Examples:
  python app.py pipeline                          Full pipeline: MoM → CRM → docs → Drive → Gmail → Slack
  python app.py pipeline --deal "Sentosa"         Pipeline for one deal
  python app.py pipeline --mom MoM.docx           Pipeline with specific MoM file
  python app.py generate                          Generate all deal documents
  python app.py generate --deal "Sentosa"         Generate for one deal
  python app.py mom MoM.docx                      Extract MoM and update CRM
  python app.py mom MoM.docx --dry-run            Preview MoM extraction only
  python app.py health                            Run deal health scoring
  python app.py interactive                       Interactive Monday.com chat
""",
    )
    sub = parser.add_subparsers(dest="command")

    # pipeline (the one-shot command)
    pipe = sub.add_parser("pipeline", help="Full pipeline: MoM → CRM update → generate docs → Drive → Gmail → Slack")
    pipe.add_argument("--mom", dest="docx_path",
                      default=os.path.join(TEMPLATE_DIR, "MoM-Sentosa-Health-PoC-Review.docx"),
                      help="Path to MoM .docx file (default: sample MoM)")
    pipe.add_argument("--deal", default="", help="Filter by deal name (substring match)")

    # generate
    gen = sub.add_parser("generate", help="Generate documents from Monday.com deals")
    gen.add_argument("--deal", default="", help="Filter by deal name (substring match)")

    # mom
    mom = sub.add_parser("mom", help="Extract MoM and update Monday.com CRM")
    mom.add_argument("docx_path", nargs="?",
                     default=os.path.join(TEMPLATE_DIR, "MoM-Sentosa-Health-PoC-Review.docx"),
                     help="Path to MoM .docx file")
    mom.add_argument("--dry-run", action="store_true", help="Preview only, don't update CRM")

    # health
    sub.add_parser("health", help="Run deal health scoring")

    # interactive
    sub.add_parser("interactive", help="Interactive Monday.com assistant")

    # serve
    srv = sub.add_parser("serve", help="Start API server for dashboard")
    srv.add_argument("--port", type=int, default=5005, help="Port (default 5005)")
    srv.add_argument("--host", default="0.0.0.0", help="Host (default 0.0.0.0)")

    args = parser.parse_args()

    if args.command == "pipeline":
        cmd_pipeline(args)
    elif args.command == "generate":
        cmd_generate(args)
    elif args.command == "mom":
        cmd_mom(args)
    elif args.command == "health":
        cmd_health(args)
    elif args.command == "interactive":
        cmd_interactive(args)
    elif args.command == "serve":
        app = create_app()
        print(f"Dashboard: http://localhost:{args.port}")
        print(f"API sync:  POST http://localhost:{args.port}/api/sync")
        print(f"API data:  GET  http://localhost:{args.port}/api/dashboard")
        app.run(host=args.host, port=args.port, debug=True)
    else:
        parser.print_help()


if __name__ == "__main__":
    main()
