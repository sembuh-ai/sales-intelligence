#!/usr/bin/env python3
"""CLI tool — Extract MoM data from .docx and update existing Monday.com CRM deals."""

import asyncio
import argparse
import json
import os
import sys
from contextlib import AsyncExitStack

import requests
from anthropic import Anthropic
from docx import Document
from dotenv import load_dotenv
from mcp import ClientSession, StdioServerParameters
from mcp.client.stdio import stdio_client

load_dotenv(os.path.join(os.path.dirname(os.path.abspath(__file__)), ".env"))

CLAUDE_MODEL = os.getenv("CLAUDE_MODEL", "claude-sonnet-4-5-20250929")
MONDAY_API_KEY = os.getenv("MONDAY_API_KEY", "")
MONDAY_WORKSPACE_ID = os.getenv("MONDAY_WORKSPACE_ID", "")
MONDAY_API_URL = "https://api.monday.com/v2"
DEALS_BOARD_ID = 5027997854

SYSTEM_PROMPT = f"""\
You are a Sales Intelligence CRM updater.
WORKSPACE: Only operate within workspace ID {MONDAY_WORKSPACE_ID}.
CRITICAL: NEVER create new items/deals. Only search and add comments on EXISTING items.
"""


def extract_docx_text(filepath: str) -> str:
    """Extract all text and tables from a .docx file."""
    doc = Document(filepath)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for i, table in enumerate(doc.tables):
        parts.append(f"\n=== TABLE {i + 1} ===")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append(" | ".join(cells))
    return "\n".join(parts)


def monday_api(query: str, variables: dict = None) -> dict:
    """Call Monday.com GraphQL API directly."""
    headers = {
        "Authorization": MONDAY_API_KEY,
        "Content-Type": "application/json",
        "API-Version": "2024-10",
    }
    payload = {"query": query}
    if variables:
        payload["variables"] = variables
    resp = requests.post(MONDAY_API_URL, json=payload, headers=headers)
    data = resp.json()
    if "errors" in data:
        raise RuntimeError(f"Monday API error: {data['errors']}")
    return data


def monday_search_deal(deal_name: str) -> list:
    """Search for deals by name on the Deals board."""
    query = f"""
    query {{
        boards(ids: [{DEALS_BOARD_ID}]) {{
            items_page(limit: 10, query_params: {{rules: [{{column_id: "name", compare_value: ["{deal_name}"], operator: contains_text}}]}}) {{
                items {{
                    id
                    name
                    column_values {{
                        id
                        text
                        value
                    }}
                }}
            }}
        }}
    }}
    """
    result = monday_api(query)
    items = result.get("data", {}).get("boards", [{}])[0].get("items_page", {}).get("items", [])
    return items


def monday_update_columns(item_id: int, column_values: dict) -> dict:
    """Update column values on an existing Monday.com item."""
    query = """
    mutation ($boardId: ID!, $itemId: ID!, $columnValues: JSON!) {
        change_multiple_column_values(
            board_id: $boardId,
            item_id: $itemId,
            column_values: $columnValues
        ) {
            id
            name
        }
    }
    """
    variables = {
        "boardId": str(DEALS_BOARD_ID),
        "itemId": str(item_id),
        "columnValues": json.dumps(column_values),
    }
    return monday_api(query, variables)


def monday_add_comment(item_id: int, text: str) -> dict:
    """Add a comment/update on a Monday.com item."""
    query = """
    mutation ($itemId: ID!, $body: String!) {
        create_update(item_id: $itemId, body: $body) {
            id
        }
    }
    """
    return monday_api(query, {"itemId": str(item_id), "body": text})


async def run(docx_path: str, dry_run: bool = False):
    """Extract MoM and update existing Monday.com CRM deal."""
    # 1. Extract MoM
    print(f"Extracting MoM from: {docx_path}")
    mom_text = extract_docx_text(docx_path)
    print(f"Extracted {len(mom_text)} chars.\n")

    # 2. Use Claude to extract structured CRM data
    print("Analyzing MoM with Claude...")
    anthropic = Anthropic()
    extract_resp = anthropic.messages.create(
        model=CLAUDE_MODEL,
        max_tokens=1024,
        messages=[{"role": "user", "content": f"""\
Analyze this Minutes of Meeting and extract structured CRM data as JSON.

--- MoM ---
{mom_text}
--- End ---

Return ONLY valid JSON (no markdown fences, no explanation):
{{
  "deal_name": "client/deal name to search in CRM",
  "deal_stage": "one of: Waiting Confirmation, Won, Lost, Solutioning, First Meeting Done, Piloting",
  "deal_value": number or null,
  "last_interaction_date": "YYYY-MM-DD or null",
  "expected_close_date": "YYYY-MM-DD or null",
  "close_probability": number 0-100 or null,
  "forecast_category": "one of: Best case, Pipeline, Commit, or null",
  "comment": "concise MoM summary under 400 chars plain ASCII. Cover: date, attendees, summary, key actions, next steps, blockers"
}}
"""}],
    )
    raw_json = extract_resp.content[0].text.strip()
    if raw_json.startswith("```"):
        raw_json = raw_json.split("\n", 1)[1].rsplit("```", 1)[0].strip()
    crm_data = json.loads(raw_json)
    print(f"\nExtracted CRM data:")
    print(json.dumps(crm_data, indent=2))

    if dry_run:
        print("\n=== DRY RUN — No Monday.com update performed ===")
        return

    # 3. Search for existing deal
    print(f"\nSearching for deal: {crm_data['deal_name']}...")
    items = monday_search_deal(crm_data["deal_name"])

    if not items:
        print(f"No deal found matching '{crm_data['deal_name']}'. Aborting (no new deals created).")
        return

    # Display matches
    print(f"Found {len(items)} matching deal(s):")
    for item in items:
        cols = {cv["id"]: cv["text"] for cv in item["column_values"]}
        print(f"  [{item['id']}] {item['name']} — Stage: {cols.get('deal_stage', '?')}, Value: {cols.get('deal_value', '?')}")

    # Pick best match (prefer exact name match with STP reference from MoM)
    target = items[0]
    for item in items:
        if "STP" in item["name"]:
            target = item
            break

    item_id = int(target["id"])
    print(f"\nUpdating deal: [{item_id}] {target['name']}")

    # 4. Update columns via Monday API
    column_values = {}
    if crm_data.get("deal_stage"):
        column_values["deal_stage"] = {"label": crm_data["deal_stage"]}
    if crm_data.get("deal_value") is not None:
        column_values["deal_value"] = str(crm_data["deal_value"])
    if crm_data.get("last_interaction_date"):
        column_values["date__1"] = {"date": crm_data["last_interaction_date"]}
    if crm_data.get("close_probability") is not None:
        column_values["deal_close_probability"] = str(crm_data["close_probability"])
    if crm_data.get("expected_close_date"):
        column_values["deal_expected_close_date"] = {"date": crm_data["expected_close_date"]}
    if crm_data.get("forecast_category"):
        column_values["color_mm1hjwv8"] = {"label": crm_data["forecast_category"]}

    if column_values:
        print(f"Updating columns: {list(column_values.keys())}")
        try:
            monday_update_columns(item_id, column_values)
            print("Columns updated successfully.")
        except Exception as e:
            print(f"Column update error: {e}", file=sys.stderr)
    else:
        print("No column values to update.")

    # 5. Add MoM comment
    if crm_data.get("comment"):
        print("Adding MoM comment...")
        try:
            monday_add_comment(item_id, crm_data["comment"])
            print("Comment added successfully.")
        except Exception as e:
            print(f"Comment error: {e}", file=sys.stderr)

    print(f"\nDone. Existing deal [{item_id}] {target['name']} updated.")


def main():
    parser = argparse.ArgumentParser(
        description="Extract MoM from .docx and update existing Monday.com CRM deal"
    )
    script_dir = os.path.dirname(os.path.abspath(__file__))
    parser.add_argument(
        "docx_path",
        nargs="?",
        default=os.path.join(
            script_dir,
            "PTP Hackathon - Brief",
            "MoM-Sentosa-Health-PoC-Review.docx",
        ),
        help="Path to MoM .docx file (default: MoM-Sentosa-Health-PoC-Review.docx)",
    )
    parser.add_argument(
        "--dry-run",
        action="store_true",
        help="Extract MoM only, don't update Monday.com",
    )
    args = parser.parse_args()

    if not os.path.exists(args.docx_path):
        print(f"File not found: {args.docx_path}", file=sys.stderr)
        sys.exit(1)

    asyncio.run(run(args.docx_path, args.dry_run))


if __name__ == "__main__":
    main()
